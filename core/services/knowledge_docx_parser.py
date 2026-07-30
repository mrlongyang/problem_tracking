import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from io import BytesIO
from typing import Iterator

from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


ERROR_PATTERN = re.compile(
    r"""
    (?P<code>[A-Z][A-Z0-9_-]{1,30})
    \s*[:：]\s*
    \[
        (?P<title>.+?)
    \]
    """,
    re.VERBOSE,
)

FUNCTION_PATTERN = re.compile(
    r"""
    ^\s*
    (?:\d+(?:\.\d+)*\.?\s*)?
    (?P<name>.+?)
    \s*
    \[
        (?P<transaction>\d{5,10})
    \]
    \s*$
    """,
    re.VERBOSE,
)

MODULE_PATTERN = re.compile(
    r"Function\s+Module\s*[（(]\s*(?P<module>[^）)]+)\s*[）)]",
    re.IGNORECASE,
)

@dataclass
class ParsedImage:
    file_name: str
    content_type: str
    data: bytes
    display_order: int = 0

@dataclass
class ParsedArticle:
    error_code: str
    title: str
    function_name: str = ""
    transaction_code: str = ""
    module_name: str = ""
    root_cause: str = ""
    resolution: str = ""
    source_order: int = 0
    warnings: list[str] = field(default_factory=list)
    images: list[ParsedImage] = field(default_factory=list)


def normalize_text(value: str) -> str:
    value = value or ""

    replacements = {
        "\u00a0": " ",
        "：": ":",
        "（": "(",
        "）": ")",
        "\r": "\n",
    }

    for old, new in replacements.items():
        value = value.replace(old, new)

    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)

    return value.strip()


def iter_document_text(document: Document):
    """
    Read both normal paragraphs and text inside tables.
    """

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)

        if text:
            yield text

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = normalize_text(cell.text)

                if text:
                    yield text


def append_text(
    existing_text: str,
    new_text: str,
) -> str:
    if not existing_text:
        return new_text.strip()

    return f"{existing_text}\n{new_text.strip()}"


def finalize_article(
    article: ParsedArticle,
) -> None:
    article.root_cause = normalize_text(
        article.root_cause
    )

    article.resolution = normalize_text(
        article.resolution
    )

    if not article.root_cause:
        article.warnings.append(
            "Root cause was not detected."
        )

    if not article.resolution:
        article.warnings.append(
            "Resolution was not detected."
        )

    if not article.images:
        article.warnings.append(
            "No screenshot was detected."
        )

def parse_docx(
    file_path: str | Path,
) -> list[ParsedArticle]:
    document = Document(str(file_path))

    articles: list[ParsedArticle] = []

    current_module = ""
    current_function = ""
    current_transaction = ""

    current_article: ParsedArticle | None = None
    current_section: str | None = None

    source_order = 0

    for item_type, item_value in iter_document_items(
        document
    ):
        source_order += 1

        # Attach an image to the current error article
        if item_type == "image":
            if current_article:
                image = item_value
                image.display_order = len(
                    current_article.images
                )

                current_article.images.append(image)

            continue

        line = normalize_text(str(item_value))

        if not line:
            continue

        module_match = MODULE_PATTERN.search(line)

        if module_match:
            current_module = normalize_text(
                module_match.group("module")
            )

        function_match = FUNCTION_PATTERN.match(line)

        if function_match and not ERROR_PATTERN.search(line):
            current_function = normalize_text(
                function_match.group("name")
            )

            current_transaction = normalize_text(
                function_match.group("transaction")
            )

            continue

        error_match = ERROR_PATTERN.search(line)

        if error_match:
            if current_article:
                finalize_article(current_article)
                articles.append(current_article)

            current_article = ParsedArticle(
                error_code=normalize_text(
                    error_match.group("code")
                ),
                title=normalize_text(
                    error_match.group("title")
                ),
                function_name=current_function,
                transaction_code=current_transaction,
                module_name=current_module,
                source_order=source_order,
            )

            current_section = None
            continue

        lowered = line.lower().strip()

        if re.match(
            r"^root\s*cause\s*:?",
            lowered,
            flags=re.IGNORECASE,
        ):
            current_section = "root_cause"

            content = re.sub(
                r"^root\s*cause\s*:?\s*",
                "",
                line,
                flags=re.IGNORECASE,
            )

            if current_article and content:
                current_article.root_cause = content

            continue

        if re.match(
            r"^resolution\s*:?",
            lowered,
            flags=re.IGNORECASE,
        ):
            current_section = "resolution"

            content = re.sub(
                r"^resolution\s*:?\s*",
                "",
                line,
                flags=re.IGNORECASE,
            )

            if current_article and content:
                current_article.resolution = content

            continue

        if not current_article:
            continue

        if current_section == "root_cause":
            current_article.root_cause = append_text(
                current_article.root_cause,
                line,
            )

        elif current_section == "resolution":
            current_article.resolution = append_text(
                current_article.resolution,
                line,
            )

    if current_article:
        finalize_article(current_article)
        articles.append(current_article)

    return articles

def iter_block_items(parent) -> Iterator[Paragraph | Table]:
    """
    Yield paragraphs and tables in their original DOCX order.
    """

    if isinstance(parent, DocumentObject):
        parent_element = parent.element.body

    elif isinstance(parent, _Cell):
        parent_element = parent._tc

    else:
        raise TypeError(
            f"Unsupported parent type: {type(parent)}"
        )

    for child in parent_element.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)

        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)
            
            
def extract_images_from_paragraph(
    paragraph: Paragraph,
) -> list[ParsedImage]:
    images: list[ParsedImage] = []

    blips = paragraph._element.xpath(
        ".//a:blip"
    )

    for index, blip in enumerate(blips):
        relationship_id = blip.get(
            qn("r:embed")
        )

        if not relationship_id:
            continue

        image_part = paragraph.part.related_parts.get(
            relationship_id
        )

        if not image_part:
            continue

        original_name = Path(
            str(image_part.partname)
        ).name

        extension = Path(original_name).suffix

        if not extension:
            extension = ".png"

        images.append(
            ParsedImage(
                file_name=original_name,
                content_type=image_part.content_type,
                data=image_part.blob,
                display_order=index,
            )
        )

    return images

def iter_document_items(document: DocumentObject):
    """
    Yield Word text and images in document order.

    Result format:
    ("text", "some paragraph")
    ("image", ParsedImage(...))
    """

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            text = normalize_text(block.text)

            if text:
                yield "text", text

            for image in extract_images_from_paragraph(
                block
            ):
                yield "image", image

        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for cell_block in iter_block_items(cell):
                        if not isinstance(
                            cell_block,
                            Paragraph,
                        ):
                            continue

                        text = normalize_text(
                            cell_block.text
                        )

                        if text:
                            yield "text", text

                        for image in extract_images_from_paragraph(
                            cell_block
                        ):
                            yield "image", image